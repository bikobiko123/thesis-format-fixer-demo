import argparse
import json
import re
import zipfile
from collections import Counter
from pathlib import Path
from typing import Any
from xml.etree import ElementTree as ET

from docx import Document

DEFAULT_OUTPUT_DIR = Path("samples/realistic_swufe")

BLANK_PNG = bytes.fromhex(
    "89504e470d0a1a0a0000000d4948445200000001000000010806000000"
    "1f15c4890000000a49444154789c6360000002000100ffff030000060005"
    "57bfab0000000049454e44ae426082"
)
BLANK_JPEG = BLANK_PNG
BLANK_SVG = b'<svg xmlns="http://www.w3.org/2000/svg" width="1" height="1"></svg>'


def _local_name(tag: str) -> str:
    return tag.rsplit("}", 1)[-1] if "}" in tag else tag


def _media_placeholder(name: str) -> bytes:
    suffix = Path(name).suffix.lower()
    if suffix == ".svg":
        return BLANK_SVG
    if suffix in {".jpg", ".jpeg"}:
        return BLANK_JPEG
    return BLANK_PNG


def _sanitize_xml(data: bytes, prefix: str, anonymize_all_text: bool = False) -> bytes:
    try:
        root = ET.fromstring(data)
    except ET.ParseError:
        return data

    counter = 0
    for element in root.iter():
        name = _local_name(element.tag)
        if name == "Relationship" and element.attrib.get("TargetMode") == "External":
            element.set("Target", "https://example.invalid")

        if anonymize_all_text and element.text and element.text.strip():
            counter += 1
            element.text = f"{prefix}-META-{counter:04d}"
            continue

        if name in {"t", "delText"} and element.text and element.text.strip():
            counter += 1
            element.text = f"{prefix}-P{counter:04d}"
        elif name == "instrText" and element.text and element.text.strip():
            instruction = element.text.strip().upper()
            if instruction.startswith(("PAGE", "NUMPAGES", "SECTIONPAGES", "TOC")):
                element.text = element.text
            else:
                counter += 1
                element.text = f"{prefix}-FIELD-{counter:04d}"

    return ET.tostring(root, encoding="utf-8", xml_declaration=True)


def anonymize_docx(source: Path, target: Path, prefix: str) -> Path:
    """Create a DOCX that keeps formatting structure but removes real text/media."""

    target.parent.mkdir(parents=True, exist_ok=True)
    with zipfile.ZipFile(source, "r") as src, zipfile.ZipFile(
        target, "w", compression=zipfile.ZIP_DEFLATED
    ) as dst:
        for info in src.infolist():
            name = info.filename
            if name.endswith("/"):
                dst.writestr(info, b"")
                continue

            data = src.read(name)
            if name.startswith("word/media/"):
                dst.writestr(info, _media_placeholder(name))
                continue

            if name.endswith(".xml") or name.endswith(".rels"):
                anonymize_all = name.startswith(("docProps/", "customXml/"))
                data = _sanitize_xml(data, prefix=prefix, anonymize_all_text=anonymize_all)
            dst.writestr(info, data)
    return target


def summarize_docx_format(path: Path) -> dict[str, Any]:
    doc = Document(path)
    style_counts = Counter(
        paragraph.style.name if paragraph.style else "None" for paragraph in doc.paragraphs
    )
    alignments = Counter(str(paragraph.alignment) for paragraph in doc.paragraphs)

    sections = []
    for section in doc.sections:
        sections.append(
            {
                "top_cm": round(section.top_margin.cm, 2),
                "bottom_cm": round(section.bottom_margin.cm, 2),
                "left_cm": round(section.left_margin.cm, 2),
                "right_cm": round(section.right_margin.cm, 2),
                "header_cm": round(section.header_distance.cm, 2),
                "footer_cm": round(section.footer_distance.cm, 2),
            }
        )

    run_signatures = Counter()
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            if not run.text.strip():
                continue
            run_signatures[
                (
                    run.font.name or "inherited",
                    round(run.font.size.pt, 2) if run.font.size else "inherited",
                    run.bold,
                )
            ] += 1
            break

    return {
        "file": path.name,
        "paragraph_count": len(doc.paragraphs),
        "section_count": len(doc.sections),
        "table_count": len(doc.tables),
        "sections": sections,
        "top_styles": style_counts.most_common(16),
        "alignments": alignments.most_common(),
        "top_run_signatures": [(str(key), count) for key, count in run_signatures.most_common(16)],
        "headers_nonempty": sum(
            1
            for section in doc.sections
            for paragraph in section.header.paragraphs
            if paragraph.text.strip()
        ),
        "footers_nonempty": sum(
            1
            for section in doc.sections
            for paragraph in section.footer.paragraphs
            if paragraph.text.strip()
        ),
    }


def build_delta(before_docx: Path, after_docx: Path) -> dict[str, Any]:
    before = summarize_docx_format(before_docx)
    after = summarize_docx_format(after_docx)
    before_margin = before["sections"][0] if before["sections"] else {}
    after_margin = after["sections"][0] if after["sections"] else {}
    return {
        "privacy_note": (
            "Generated from local real before/after files after replacing text nodes, "
            "document properties, custom XML text, external links, and media binaries."
        ),
        "source_policy": "附件8西南财经大学关于研究生学位论文形式与格式的基本要求.doc",
        "before": before,
        "after": after,
        "observed_format_delta": {
            "section_count": {"before": before["section_count"], "after": after["section_count"]},
            "table_count": {"before": before["table_count"], "after": after["table_count"]},
            "first_section_margins": {"before": before_margin, "after": after_margin},
            "headers_nonempty": {
                "before": before["headers_nonempty"],
                "after": after["headers_nonempty"],
            },
            "new_after_styles": sorted(
                {
                    style
                    for style, _ in after["top_styles"]
                    if style not in {before_style for before_style, _ in before["top_styles"]}
                }
            ),
        },
    }


def _assert_no_source_terms(path: Path) -> None:
    risky_terms = ["SOURCE-PRIVATE-TITLE", "SOURCE-STUDENT-ID", "SOURCE-PRIVATE-BODY"]
    with zipfile.ZipFile(path) as archive:
        payload = "\n".join(
            archive.read(name).decode("utf-8", errors="ignore")
            for name in archive.namelist()
            if re.search(r"\.(xml|rels)$", name)
        )
    for term in risky_terms:
        if term in payload:
            raise ValueError(f"Anonymization check failed for term: {term}")


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(
        description="Build anonymized SWUFE demo DOCX files from local real before/after files."
    )
    parser.add_argument("--before", required=True, type=Path, help="Local real before DOCX path.")
    parser.add_argument("--after", required=True, type=Path, help="Local real after DOCX path.")
    parser.add_argument(
        "--policy",
        required=True,
        type=Path,
        help="Local SWUFE policy .doc path, used only for source naming in metadata.",
    )
    parser.add_argument("--output-dir", default=DEFAULT_OUTPUT_DIR, type=Path)
    return parser.parse_args()


def main() -> None:
    args = parse_args()
    args.output_dir.mkdir(parents=True, exist_ok=True)
    before_out = args.output_dir / "swufe_before_anonymized.docx"
    after_out = args.output_dir / "swufe_after_anonymized.docx"
    anonymize_docx(args.before, before_out, prefix="SWUFE-BEFORE")
    anonymize_docx(args.after, after_out, prefix="SWUFE-AFTER")
    _assert_no_source_terms(before_out)
    _assert_no_source_terms(after_out)

    delta = build_delta(before_out, after_out)
    delta["source_policy"] = args.policy.name
    (args.output_dir / "format_delta.json").write_text(
        json.dumps(delta, ensure_ascii=False, indent=2),
        encoding="utf-8",
    )
    (args.output_dir / "policy_rules_used.json").write_text(
        json.dumps(
            {
                "school": "西南财经大学",
                "profile": "swufe_master",
                "source": args.policy.name,
                "rules_used": {
                    "paper": "A4",
                    "body_font": "宋体，小4号",
                    "body_line_spacing": "多倍行距 1.37",
                    "margins_cm": {"top": 3.9, "bottom": 3.4, "left": 3.45, "right": 3.45},
                    "header_footer_cm": {"header": 2.8, "footer": 2.5},
                    "heading_1": "2号小标宋",
                    "heading_2": "小3号黑体",
                    "heading_3": "小4号黑体",
                    "caption": "图表字号5号，表题黑体，内容宋体；表注小5号宋体",
                },
            },
            ensure_ascii=False,
            indent=2,
        ),
        encoding="utf-8",
    )


if __name__ == "__main__":
    main()
