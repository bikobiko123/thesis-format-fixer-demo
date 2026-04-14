from pathlib import Path

from docx import Document


def main() -> None:
    output_dir = Path("samples/input")
    output_dir.mkdir(parents=True, exist_ok=True)
    document = Document()
    document.add_paragraph("高校毕业论文格式智能修复系统研究")
    document.add_paragraph("摘要").style = "Heading 1"
    document.add_paragraph("本文提出一个面向论文提交前格式修复的 MVP 系统。")
    document.add_paragraph("关键词：毕业论文；格式修复；规则引擎")
    document.add_paragraph("Abstract").style = "Heading 1"
    document.add_paragraph("This demo repairs deterministic formatting issues in DOCX theses.")
    document.add_paragraph("目录").style = "Heading 1"
    document.add_paragraph("第一章 绪论").style = "Heading 1"
    document.add_paragraph("系统只处理文档格式，不生成论文内容。")
    document.add_paragraph("图1-1 系统处理流程")
    document.add_paragraph("参考文献").style = "Heading 1"
    document.add_paragraph("[1] Demo University. Thesis formatting guide.")
    document.save(output_dir / "demo_thesis.docx")


if __name__ == "__main__":
    main()
