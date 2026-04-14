import zipfile

from app.docx_engine.python_docx import PythonDocxEngine
from app.llm.provider import RuleBasedLLMProvider
from app.services.processor import ThesisProcessor
from docx import Document


def test_processor_returns_zip_with_repaired_docx_and_reports(tmp_path) -> None:
    source = tmp_path / "input.docx"
    document = Document()
    document.add_paragraph("高校毕业论文格式智能修复系统研究")
    document.add_paragraph("摘要").style = "Heading 1"
    document.add_paragraph("关键词：论文格式；智能修复")
    document.add_paragraph("目录").style = "Heading 1"
    document.add_paragraph("第一章 绪论").style = "Heading 1"
    document.add_paragraph("图1-1 系统架构")
    document.add_paragraph("参考文献").style = "Heading 1"
    document.save(source)

    processor = ThesisProcessor(PythonDocxEngine(), RuleBasedLLMProvider(), tmp_path / "storage")

    result = processor.process(source, "undergraduate")

    assert result.archive_path.exists()
    with zipfile.ZipFile(result.archive_path) as archive:
        assert sorted(archive.namelist()) == [
            "format_report.json",
            "format_report.md",
            "manual_fix_list.md",
            "repaired_thesis.docx",
        ]
