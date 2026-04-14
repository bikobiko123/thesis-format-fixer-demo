from pathlib import Path

from app.docx_engine.python_docx import PythonDocxEngine
from docx import Document


def test_python_docx_engine_applies_rule_style_name_and_paragraph_spacing(
    tmp_path: Path,
) -> None:
    source = tmp_path / "source.docx"
    output = tmp_path / "output.docx"
    document = Document()
    document.add_paragraph("第一章 绪论")
    document.save(source)

    PythonDocxEngine().repair(
        source,
        output,
        [
            {
                "type": "set_paragraph_style",
                "paragraph_index": 0,
                "style_key": "heading_1",
                "style": {
                    "style_name": "Heading 1",
                    "font_family": "SimHei",
                    "font_size_pt": 22,
                    "bold": False,
                    "alignment": "center",
                    "line_spacing": 1.37,
                    "first_line_indent_chars": 0,
                    "space_before_pt": 22.5,
                    "space_after_pt": 15,
                },
            }
        ],
    )

    paragraph = Document(output).paragraphs[0]

    assert paragraph.style.name == "Heading 1"
    assert str(paragraph.alignment) == "CENTER (1)"
    assert round(paragraph.paragraph_format.space_before.pt, 2) == 22.5
    assert round(paragraph.paragraph_format.space_after.pt, 2) == 15


def test_python_docx_engine_creates_missing_rule_style_name(tmp_path: Path) -> None:
    source = tmp_path / "source.docx"
    output = tmp_path / "output.docx"
    document = Document()
    document.add_paragraph("表1-1 变量说明")
    document.save(source)

    PythonDocxEngine().repair(
        source,
        output,
        [
            {
                "type": "set_paragraph_style",
                "paragraph_index": 0,
                "style_key": "caption",
                "style": {
                    "style_name": "表题西财",
                    "font_family": "Times New Roman",
                    "east_asia_font_family": "SimHei",
                    "font_size_pt": 10.5,
                    "bold": False,
                    "alignment": "center",
                    "line_spacing": 1.37,
                    "first_line_indent_chars": 0,
                    "first_line_indent_cm": 0,
                    "space_before_pt": 2.5,
                    "space_after_pt": 2.5,
                },
            }
        ],
    )

    paragraph = Document(output).paragraphs[0]

    assert paragraph.style.name == "表题西财"
    assert str(paragraph.alignment) == "CENTER (1)"
