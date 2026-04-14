import json
import subprocess
import sys
from pathlib import Path

from docx import Document


def test_cli_repair_writes_predictable_agent_artifacts(tmp_path: Path) -> None:
    source = tmp_path / "input.docx"
    output_dir = tmp_path / "agent-output"
    document = Document()
    document.add_paragraph("摘要").style = "Heading 1"
    document.add_paragraph("关键词：论文格式；智能修复")
    document.add_paragraph("目录").style = "Heading 1"
    document.add_paragraph("第一章 绪论").style = "Heading 1"
    document.add_paragraph("参考文献").style = "Heading 1"
    document.save(source)

    result = subprocess.run(
        [
            sys.executable,
            "-m",
            "app.cli",
            "repair",
            "--input",
            str(source),
            "--profile",
            "swufe_master",
            "--out",
            str(output_dir),
            "--docx-engine",
            "openxml",
            "--structure-recognizer",
            "heuristic",
        ],
        check=True,
        capture_output=True,
        text=True,
    )

    payload = json.loads(result.stdout)
    assert payload["status"] == "ok"
    assert (output_dir / "repaired_thesis.docx").exists()
    assert (output_dir / "format_report.json").exists()
    assert (output_dir / "format_report.md").exists()
    assert (output_dir / "manual_fix_list.md").exists()
    assert (output_dir / "thesis_format_fix_result.zip").exists()
