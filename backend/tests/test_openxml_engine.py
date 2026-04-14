import zipfile
from pathlib import Path

from app.docx_engine.openxml import OpenXmlDocxEngine
from app.rules.loader import load_rules
from app.structure.recognizer import HeuristicStructureRecognizer
from docx import Document


def _xml_text(path: Path, part_name: str) -> str:
    with zipfile.ZipFile(path) as archive:
        return archive.read(part_name).decode("utf-8")


def test_openxml_engine_repairs_fields_sections_page_restart_and_toc(tmp_path: Path) -> None:
    source = tmp_path / "source.docx"
    output = tmp_path / "output.docx"
    document = Document()
    document.add_paragraph("目录").style = "Heading 1"
    document.add_paragraph("第一章 绪论").style = "Heading 1"
    document.add_paragraph("正文段落")
    document.save(source)

    OpenXmlDocxEngine().repair(
        source,
        output,
        [
            {
                "type": "set_margins",
                "margins": {
                    "top_cm": 3.9,
                    "bottom_cm": 3.4,
                    "left_cm": 3.45,
                    "right_cm": 3.45,
                    "header_cm": 2.8,
                    "footer_cm": 2.5,
                },
            },
            {"type": "ensure_page_number_footer", "restart_at_body": True},
            {"type": "mark_toc_dirty"},
            {"type": "ensure_section_break_before", "paragraph_index": 1},
        ],
    )

    document_xml = _xml_text(output, "word/document.xml")
    rels_xml = _xml_text(output, "word/_rels/document.xml.rels")

    assert 'w:top="2211"' in document_xml
    assert 'w:header="1588"' in document_xml
    assert 'w:start="1"' in document_xml
    assert document_xml.count("pgNumType") >= 2
    assert 'w:val="nextPage"' in document_xml
    assert 'w:dirty="true"' in document_xml
    assert "TOC" in document_xml
    assert "footer" in rels_xml

    repaired = Document(output)
    assert len(repaired.sections) == 2
    for section in repaired.sections:
        assert round(section.top_margin.cm, 2) == 3.9
        assert round(section.bottom_margin.cm, 2) == 3.4
        assert round(section.left_margin.cm, 2) == 3.45
        assert round(section.right_margin.cm, 2) == 3.45

    repaired_ir = HeuristicStructureRecognizer().recognize(
        OpenXmlDocxEngine().parse(output),
        load_rules("swufe_master"),
    )
    assert "Table of Contents" in repaired_ir.sections

    with zipfile.ZipFile(output) as archive:
        footer_parts = [name for name in archive.namelist() if name.startswith("word/footer")]
        footer_payload = "\n".join(archive.read(name).decode("utf-8") for name in footer_parts)
    assert "PAGE" in footer_payload
