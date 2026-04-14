import zipfile
from pathlib import Path

from docx import Document

from scripts.build_anonymized_real_demo import anonymize_docx, summarize_docx_format


def test_anonymize_docx_removes_text_and_replaces_media(tmp_path: Path) -> None:
    source = tmp_path / "source.docx"
    target = tmp_path / "target.docx"
    source_phrase = "真实论文标题与敏感正文"

    document = Document()
    document.add_paragraph(source_phrase)
    document.add_paragraph("参考文献").style = "Heading 1"
    document.save(source)

    anonymize_docx(source, target, prefix="BEFORE")

    with zipfile.ZipFile(target) as archive:
        xml_payload = "\n".join(
            archive.read(name).decode("utf-8", errors="ignore")
            for name in archive.namelist()
            if name.endswith(".xml")
        )

    assert source_phrase not in xml_payload
    assert "BEFORE-P0001" in xml_payload
    assert summarize_docx_format(target)["paragraph_count"] == 2
