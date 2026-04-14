from pathlib import Path
from zipfile import ZipFile

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

from app.docx_engine.base import (
    DocxDocumentInfo,
    HeaderFooterInfo,
    ParagraphInfo,
    SectionInfo,
)

ALIGNMENT_TO_DOCX = {
    "left": WD_ALIGN_PARAGRAPH.LEFT,
    "center": WD_ALIGN_PARAGRAPH.CENTER,
    "right": WD_ALIGN_PARAGRAPH.RIGHT,
    "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
}


def _font_families(style: dict) -> tuple[str, str]:
    ascii_font = style.get("ascii_font_family") or style["font_family"]
    east_asia_font = style.get("east_asia_font_family") or style["font_family"]
    return ascii_font, east_asia_font


def _cm(value) -> float | None:
    return round(value.cm, 2) if value is not None else None


def _paragraph_font(paragraph) -> tuple[str | None, float | None, bool | None]:
    first_run = paragraph.runs[0] if paragraph.runs else None
    if not first_run:
        return None, None, None
    font_name = first_run.font.name
    size = round(first_run.font.size.pt, 2) if first_run.font.size else None
    return font_name, size, first_run.bold


def _toc_field_state(path: Path) -> tuple[bool, bool]:
    try:
        with ZipFile(path) as package:
            document_xml = package.read("word/document.xml").decode("utf-8", errors="ignore")
    except KeyError:
        return False, False
    return "TOC" in document_xml, "w:dirty" in document_xml and "TOC" in document_xml


def _add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = "PAGE"
    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char_begin)
    run._r.append(instr_text)
    run._r.append(fld_char_end)


def _ensure_paragraph_style(document: Document, style: dict):
    style_name = style.get("style_name")
    if not style_name:
        return None
    try:
        target_style = document.styles[style_name]
    except KeyError:
        target_style = document.styles.add_style(style_name, WD_STYLE_TYPE.PARAGRAPH)

    ascii_font, east_asia_font = _font_families(style)
    target_style.font.name = ascii_font
    target_style.font.size = Pt(style["font_size_pt"])
    target_style.font.bold = style["bold"]
    target_style._element.rPr.rFonts.set(qn("w:ascii"), ascii_font)
    target_style._element.rPr.rFonts.set(qn("w:hAnsi"), ascii_font)
    target_style._element.rPr.rFonts.set(qn("w:eastAsia"), east_asia_font)
    target_style.paragraph_format.alignment = ALIGNMENT_TO_DOCX.get(style["alignment"])
    target_style.paragraph_format.line_spacing = style["line_spacing"]
    return target_style


class PythonDocxEngine:
    """MVP DOCX engine backed by python-docx."""

    def parse(self, path: Path) -> DocxDocumentInfo:
        document = Document(path)
        paragraphs: list[ParagraphInfo] = []
        for index, paragraph in enumerate(document.paragraphs):
            font_name, size, bold = _paragraph_font(paragraph)
            paragraphs.append(
                ParagraphInfo(
                    index=index,
                    text=paragraph.text.strip(),
                    style_name=paragraph.style.name if paragraph.style else None,
                    alignment=str(paragraph.alignment) if paragraph.alignment is not None else None,
                    font_name=font_name,
                    font_size_pt=size,
                    is_bold=bold,
                )
            )

        sections: list[SectionInfo] = []
        headers: list[HeaderFooterInfo] = []
        footers: list[HeaderFooterInfo] = []
        for index, section in enumerate(document.sections):
            sections.append(
                SectionInfo(
                    index=index,
                    top_margin_cm=_cm(section.top_margin),
                    bottom_margin_cm=_cm(section.bottom_margin),
                    left_margin_cm=_cm(section.left_margin),
                    right_margin_cm=_cm(section.right_margin),
                    header_distance_cm=_cm(section.header_distance),
                    footer_distance_cm=_cm(section.footer_distance),
                )
            )
            headers.append(HeaderFooterInfo(index, "header", section.header.paragraphs[0].text))
            footers.append(HeaderFooterInfo(index, "footer", section.footer.paragraphs[0].text))

        has_toc_field, toc_needs_update = _toc_field_state(path)
        return DocxDocumentInfo(
            paragraphs=paragraphs,
            sections=sections,
            headers=headers,
            footers=footers,
            has_toc_field=has_toc_field,
            toc_needs_update=toc_needs_update,
        )

    def repair(self, source_path: Path, output_path: Path, operations: list[dict]) -> Path:
        document = Document(source_path)

        for operation in operations:
            op_type = operation.get("type")
            if op_type == "set_margins":
                margins = operation["margins"]
                for section in document.sections:
                    section.top_margin = Cm(margins["top_cm"])
                    section.bottom_margin = Cm(margins["bottom_cm"])
                    section.left_margin = Cm(margins["left_cm"])
                    section.right_margin = Cm(margins["right_cm"])
            elif op_type == "set_paragraph_style":
                index = operation["paragraph_index"]
                if index >= len(document.paragraphs):
                    continue
                paragraph = document.paragraphs[index]
                style = operation["style"]
                if target_style := _ensure_paragraph_style(document, style):
                    paragraph.style = target_style
                paragraph.alignment = ALIGNMENT_TO_DOCX.get(style["alignment"])
                fmt = paragraph.paragraph_format
                fmt.line_spacing = style["line_spacing"]
                indent_cm = style.get("first_line_indent_cm")
                if indent_cm is None:
                    indent_cm = style["first_line_indent_chars"] * 0.37
                fmt.first_line_indent = Cm(indent_cm)
                if style.get("space_before_pt") is not None:
                    fmt.space_before = Pt(style["space_before_pt"])
                if style.get("space_after_pt") is not None:
                    fmt.space_after = Pt(style["space_after_pt"])
                ascii_font, east_asia_font = _font_families(style)
                for run in paragraph.runs or [paragraph.add_run()]:
                    run.font.name = ascii_font
                    run._element.rPr.rFonts.set(qn("w:ascii"), ascii_font)
                    run._element.rPr.rFonts.set(qn("w:hAnsi"), ascii_font)
                    run._element.rPr.rFonts.set(qn("w:eastAsia"), east_asia_font)
                    run.font.size = Pt(style["font_size_pt"])
                    run.bold = style["bold"]
            elif op_type == "ensure_page_number_footer":
                for section in document.sections:
                    paragraph = section.footer.paragraphs[0]
                    if "PAGE" not in paragraph._p.xml:
                        paragraph.clear()
                        _add_page_number(paragraph)
            elif op_type == "ensure_header_text":
                text = operation["text"]
                for section in document.sections:
                    paragraph = section.header.paragraphs[0]
                    if not paragraph.text.strip():
                        paragraph.text = text

        output_path.parent.mkdir(parents=True, exist_ok=True)
        document.save(output_path)
        return output_path
